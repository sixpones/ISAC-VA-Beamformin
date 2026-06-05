"""
生成 ISAC 研究进展汇总报告 (Word 文档)
整合所有代码功能、实验结果图和论文框架
"""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT

doc = Document()

# ── 全局样式 ──────────────────────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

# ── 封面标题 ──────────────────────────────────────────────────────────────
title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_para.add_run(
    "ISAC 感知辅助自适应波束追踪\n"
    "—— 基于虚拟锚点与双域几何匹配的波束赋形方案\n"
    "代码实现与实验结果汇总报告"
)
title_run.bold = True
title_run.font.size = Pt(18)
title_run.font.color.rgb = RGBColor(0, 0, 0)

doc.add_paragraph()

info_para = doc.add_paragraph()
info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
info_para.add_run("技术栈: Sionna RT + Mitsuba + Python + Open3D + Plotly\n").font.size = Pt(10)
info_para.add_run("数据集: 8×8 MIMO UPA @ 28 GHz, METIS 2×2 变高街区场景\n").font.size = Pt(10)
info_para.add_run("日期: 2026-06-01").font.size = Pt(10)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 1. 论文概念与创新点
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading("1. 论文核心思路与创新点", level=1)

doc.add_paragraph(
    "目标: 在 5G/6G ISAC (Integrated Sensing and Communication) 架构下, "
    "利用基站感知数据（点云）辅助波束管理, 用 1 次波束扫描的代价逼近 64 次穷举扫描的频谱效率。"
)

doc.add_heading("1.1 核心流程", level=2)
steps = [
    "① ISAC 感知: 基站发射感知信号, 接收回波生成环境 3D 点云 (isac_pure_sensing.h5)",
    "② 墙面提取: 对点云做 RANSAC 平面拟合, 提取建筑物墙面方程 (wall_extraction.py)",
    "③ 虚拟锚点 (VA) 定位: 求 BS 关于墙面的镜像点, 得到 VA 三维坐标 (VAprompt.py)",
    "④ 双域几何匹配: 用 VA 坐标计算反射径的理论时延 τ 和发射角 AoD φ, "
    "在 Sionna 多径数据中匹配对应路径 (va_matching.py)",
    "⑤ 自适应波束选择: 每帧先尝试 LoS 径, VA 几何约束验证真伪, 不可用时退回反射径 "
    "(beamforming_evaluation_new.py, 方案 E)",
]
for s in steps:
    doc.add_paragraph(s, style='List Bullet')

doc.add_heading("1.2 创新点", level=2)
innovations = [
    ("创新 1 — VA 辅助双域几何匹配: ",
     "不同于传统 argmax 盲目选最强径, 利用 ISAC 感知点云提取的墙面/VA 信息, "
     "通过 3D 线面交点反算真实反射点, 在时延+角度双域锁定物理反射径, "
     "NLoS 区域匹配率达 100% (27/27 帧)。"),
    ("创新 2 — 穿透伪径鉴别: ",
     "Sionna 薄层材质模型 (thickness=0.1) 下, 射线可穿透 60m 建筑产生伪 LoS 径。"
     "传统 LoS 追踪仅靠 τ/φ 匹配, 在 NLoS 区域被穿透径欺骗 (虚高 SE)。"
     "VA 几何验证可天然筛除此类无效径 — 若 LoS 径 RSRP 高于反射径且 UE 位于 NLoS 区域, "
     "则该 LoS 径必为穿透伪径。这一物理约束机制本身就是独立贡献点。"),
    ("创新 3 — ISAC 零阶保持记忆: ",
     "VA 匹配短暂失效时 (反射点超出墙面物理边界), 沿用上一帧有效 RSRP, "
     "避免通信中断, 实现感知→通信的跨域信息融合。"),
    ("创新 4 — 自适应波束选择 (方案 E): ",
     "LoS 可用时对齐穷举搜索上界 (13.154 bps/Hz), 不可用时退回 VA 反射径 (5.796 bps/Hz), "
     "仅需 1 次波束扫描, 较 64 次穷举降低 98.4% 训练开销。"),
]
for title, desc in innovations:
    p = doc.add_paragraph()
    p.add_run(title).bold = True
    p.add_run(desc)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 2. 代码架构总览
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading("2. 代码架构与功能清单", level=1)

doc.add_paragraph(
    "整个项目包含 11 个 Python 脚本, 分为数据生成、感知特征提取、多径匹配验证、"
    "波束赋形评估四个阶段。所有代码位于 ISAC_beam/ 目录。"
)

# 2.1 数据生成
doc.add_heading("2.1 阶段一: 数据生成", level=2)

doc.add_heading("ue_multipath_sim.py — MIMO 信道数据生成", level=3)
doc.add_paragraph(
    "基于 Sionna RT (Mitsuba CUDA 后端) 生成 8×8 MIMO 射线追踪多径信道数据。\n"
    "场景: 3 栋变高建筑 (25m/30m/35m), 60×60m 底面, METIS 街区布局。\n"
    "BS: 8×8 UPA @ [40, -40, 15], 朝向 135°。\n"
    "UE: 单天线, 沿 X=0, Y=[-20, 40] 移动, 共 122 帧。\n"
    "射线追踪: max_depth=2, diffuse_reflection=True, samples=200k。\n"
    "材质: custom_concrete (εr=5.0, σ=0.05, scattering=0.7, thickness=0.1)。\n"
    "输出: ue_mimo_multipath_data.npz (607 MB) — 含 ue_positions, taus, phi_ts, h_complexs。"
)

doc.add_heading("dataset_gen_sensing_v0131.py — ISAC 感知点云生成", level=3)
doc.add_paragraph(
    "生成 ISAC 雷达感知点云数据。\n"
    "BS 发射→场景反射→BS 接收 (单站雷达模式)。\n"
    "从 τ·c/2 和到达角反推散射点 3D 坐标。\n"
    "输出: isac_pure_sensing.h5 (5.1 MB) — shape (N, 4096, 5), "
    "每点 [x, y, z, bs_index, path_gain]。\n"
    "注意: 此脚本的重心是「感知数据生成」而非「通信信道生成」, "
    "与 ue_multipath_sim.py 互补。"
)

# 2.2 感知特征提取
doc.add_heading("2.2 阶段二: ISAC 感知特征提取", level=2)

doc.add_heading("wall_extraction.py / wall_extraction_single.py — RANSAC 墙面提取", level=3)
doc.add_paragraph(
    "从 ISAC 感知点云中提取建筑物墙面:\n"
    "① 加载 isac_pure_sensing.h5, 按 path_gain 过滤无效点。\n"
    "② Open3D RANSAC 迭代平面拟合 (distance_threshold=0.5, num_iterations=1000)。\n"
    "③ 地面过滤: 法向量 Z 分量 > 0.9 则剔除。\n"
    "④ 小平面过滤: 内点 < 300 则剔除。\n"
    "⑤ Plotly 3D 可视化输出 (wall_extraction.html / wall_extraction_single.html)。\n"
    "wall_extraction_single.py 额外按 BS ID 拆分了单基站视角。"
)

doc.add_heading("VAprompt.py — 虚拟锚点计算", level=3)
doc.add_paragraph(
    "核心功能: 将 RANSAC 墙面方程转化为虚拟锚点三维坐标。\n"
    "数学原理: VA = BS 关于墙面的镜像点, 公式 p' = p - 2·(n·p + d)·n。\n"
    "物理约束剪枝:\n"
    "  ① 朝向过滤: 若 v·n < 0 (法向量背离基站), 翻转法向量方向。\n"
    "  ② BS-墙面映射: BS_0 仅对应 Wall 0 和 Wall 1。\n"
    "输出: virtual_anchor.json — 两个 VA 的三维坐标。"
)

doc.add_heading("dataset_test_sensing_v0131.py — 感知点云可视化", level=3)
doc.add_paragraph(
    "读取 isac_pure_sensing.h5, 用 Plotly Mesh3d 绘制建筑立方体和散射点云。\n"
    "输出: view_metis_v35.html (4.8 MB) — 可交互的 3D 场景视图。"
)

doc.add_page_break()

# 2.3 多径匹配验证
doc.add_heading("2.3 阶段三: 多径验证与 VA 匹配", level=2)

doc.add_heading("plot_multipath.py — MIMO 数据集全径验证", level=3)
doc.add_paragraph(
    "对 ue_mimo_multipath_data.npz 做全面物理验证:\n"
    "• 面板 1 (增益): 灰色散点为全径分量, 蓝线为 argmax 最强径, "
    "蓝虚线为 LoS 几何追踪, 红虚线为东墙 VA 反射追踪。\n"
    "• 面板 2 (ToA): 黑色虚线为理论 LoS 时延, 验证射线追踪精度。\n"
    "• 面板 3 (AoD): 黑色虚线为理论 LoS 发射角。\n"
    "关键发现: 在 NLoS 区域 (Y>26.7), argmax 仍然追踪到与 LoS 时延/角度一致的路径, "
    "但增益比真实反射径高约 15 dB — 证明这是穿透建筑物的伪径。\n"
    "输出: mimo_dataset_raw_verification.png (1.1 MB)。"
)

doc.add_heading("plot_multipath_sanity.py — 非 MIMO 数据集验证", level=3)
doc.add_paragraph(
    "对 ue_multipath_data.npz (非 MIMO, 标量增益) 做类似的 3 面板验证。\n"
    "散点背景做下采样 (80 径/帧) 以保持可读性。\n"
    "输出: multipath_sanity_check.png (1.7 MB)。"
)

doc.add_heading("va_matching.py — 虚拟锚点辅助多径几何追踪 (核心算法)", level=3)
doc.add_paragraph(
    "这是整个项目的核心匹配算法, 实现三条物理路径的显式追踪:\n\n"
    "【双域匹配机制】\n"
    "对每条路径独立计算理论时延 τ_th 和发射角 φ_th, 在 Sionna 多径集合中搜索满足条件的径:\n"
    "  · 时延条件: |τ_meas - τ_th| < 1 ns (对应 ~0.3m 距离分辨率)\n"
    "  · 角度条件: |phase_diff(φ_meas, φ_th)| < 0.1 rad (~5.7°)\n"
    "  · 多候选时选 L2 范数最大者\n\n"
    "【三条被追踪路径】\n"
    "  ① LoS 直射径: τ = |UE - BS|/C, φ = atan2(UE_y - BS_y, UE_x - BS_x)\n"
    "  ② 东墙 VA 反射: τ = |UE - VA_EAST|/C, "
    "反射点 R = 线面交点(VA_EAST→UE 与 东墙平面), φ = BS→R 方向角\n"
    "  ③ 南墙 VA 反射: 同理, 使用 VA_SOUTH 和南墙平面\n\n"
    "【3D 线面交点 (reflection_point 函数)】\n"
    "射线参数方程 P(t) = VA + t·(UE - VA) 代入墙面平面方程 n·(P - p0) = 0:\n"
    "  t = n·(p0 - VA) / n·(UE - VA)\n"
    "  R = VA + t·(UE - VA)\n"
    "VA 仅用于算时延 (镜像等价性), AoD 必须用真实的反射点 R 来算。\n\n"
    "【关键结论】\n"
    "  · LoS 区域 (95 帧): 东墙 VA 匹配 77/95 (81%), 缺失的 18 帧因反射点超出墙面物理边界\n"
    "  · NLoS 区域 (27 帧): 东墙 VA 匹配 27/27 (100%), 证明 VA 方法在 NLoS 区域高度可靠\n"
    "  · 南墙 VA: LoS 区域匹配较少, NLoS 区域 0/27 (南墙在 BS 背后, 反射径被遮挡)\n"
    "  · argmax 被穿透径劫持: NLoS 区域中 argmax 追踪到的径与 LoS 理论 τ 一致,\n"
    "    但增益比东墙真实反射径高约 15 dB — 这是薄层材质模型的已知伪影\n\n"
    "输出: va_matching_result.png (283 KB) + va_matching_diagnostics.png (664 KB)。"
)

doc.add_page_break()

# 2.4 波束赋形评估
doc.add_heading("2.4 阶段四: 波束赋形频谱效率评估", level=2)

doc.add_heading("beamforming_evaluation.py — 3 方案初版", level=3)
doc.add_paragraph(
    "初版波束赋形评估, 使用非 MIMO 数据集 (ue_multipath_data.npz), 标量增益。\n"
    "3 方案: Exhaustive, LoS-only, VA-assisted。\n"
    "输出: beamforming_evaluation.png (309 KB)。"
)

doc.add_heading("beamforming_evaluation_new.py — 5 方案完整版 (当前主实验)", level=3)
doc.add_paragraph(
    "完整波束赋形频谱效率评估, 使用 MIMO 数据集 (ue_mimo_multipath_data.npz), "
    "8×8 UPA 复信道。\n\n"
    "【五种对比方案】\n"
    "  A — MRT 穷举搜索 (理想上限): 逐径 L2 范数, 取全局最强, MRT 波束赋形。64 次扫描。\n"
    "  B — 3GPP 2D-DFT 码本穷举 (新基线): 生成符合 NR Type I 的 64 个 DFT 波束, "
    "对全径合成信道 H_total 遍历搜索。64 次扫描。\n"
    "  C — LoS 直射径追踪 (传统基线): 时延+角度双域匹配 LoS 理论值, MRT 发射。1 次扫描。\n"
    "  D — VA 纯反射 (消融分析): 仅用东墙 VA 反射径, 零阶保持记忆。1 次扫描。\n"
    "  E — VA 自适应波束 (提出的完整方案, 待加入几何验证约束后): "
    "LoS 可用且 RSRP 优于 VA 时选 LoS, 否则退回 VA。1 次扫描。\n\n"
    "【新增技术组件】\n"
    "  · generate_2d_dft_codebook(Ny=8, Nz=8): Kronecker 积生成 64 个归一化 DFT 波束向量。\n"
    "  · compute_h_total(hh, ta): Σ h_k · exp(-j2πfcτ_k) 全径合成等效信道。\n"
    "  · compute_mrt_rsrp(h_k, tau_k): MRT 权值 w = h_k^H/||h_k||, RSRP = |h·w·exp(-j2πfcτ)|²。\n\n"
    "【实验结果 (bps/Hz)】\n"
    "  A: LoS=13.154  NLoS=5.796  全局=11.525  (64 次扫描)\n"
    "  B: LoS=12.606  NLoS=5.949  全局=11.133  (64 次扫描, 量化损耗 ~0.4 bps/Hz)\n"
    "  C: LoS=13.154  NLoS=5.796  全局=11.525  (1 次扫描, 但 NLoS 被穿透伪径欺骗)\n"
    "  D: LoS=8.636   NLoS=5.796  全局=8.007   (1 次扫描, 纯反射径)\n"
    "  E: LoS=13.154  NLoS=5.796  全局=11.525  (1 次扫描, 当前退化为 C, 需加入几何约束)\n\n"
    "【待解决问题】\n"
    "方案 E 目前的 'if rsrp_los > rsrp_va' 比较在 NLoS 区域也被穿透伪径满足, "
    "导致退化为方案 C。需要加入几何位置约束 (y_coord < LOS_BLOCKAGE_Y) "
    "来实现真正的'VA 鉴别穿透伪径'。\n\n"
    "输出: isac_beamforming_eval_5schemes.png (516 KB)。"
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 3. 实验结果图汇总
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading("3. 实验结果图汇总", level=1)

beam_dir = "/home/xu/xdh/isac/ISAC_beam"

figures = [
    ("Fig. 1: 场景几何与感知点云", "view_metis_v35.html",
     "METIS 2×2 街区 3D 可视化 (Plotly HTML 交互)。展示 3 栋变高建筑、3 个 BS 位置 (红钻) "
     "及 ISAC 感知点云覆盖范围。可在浏览器中旋转/缩放截图用于论文。",
     "论文位置建议: Section II System Model / 场景示意图"),

    ("Fig. 2: RANSAC 墙面提取", "wall_extraction_single.html",
     "从 ISAC 感知点云中通过 RANSAC 迭代平面拟合提取的建筑物墙面 (蓝色/绿色点), "
     "灰色为离群点, 红色菱形为 BS_0。这是 VA 定位的前置步骤。",
     "论文位置建议: Section III ISAC Sensing / 感知辅助特征提取"),

    ("Fig. 3: VA 定位与可视化", "va_visualization.html",
     "RANSAC 墙面 (彩色) + 基站 (红色菱形) + 虚拟锚点 (蓝色圆圈) + BS↔VA 虚线连线。"
     "展示 VA 与真实墙面的镜像几何关系。",
     "论文位置建议: Section III ISAC Sensing / VA 定位原理"),

    ("Fig. 4: 多径数据集验证 (MIMO)", "mimo_dataset_raw_verification.png",
     "3 面板验证图: (上) 全径散点 + argmax 最强径 + LoS/VA 几何追踪增益对比; "
     "(中) ToA 对比, 黑色虚线为理论 LoS; (下) AoD 对比。"
     "关键发现: NLoS 区域 argmax 仍追踪到与 LoS 一致的 τ/φ, 但增益低 15 dB → 穿透伪径。",
     "论文位置建议: Appendix A / 数据集物理验证"),

    ("Fig. 5: VA 双域匹配诊断 (核心)", "va_matching_diagnostics.png",
     "3 面板诊断图: (上) 增益 — 灰虚线 argmax、蓝虚线 LoS、红实线 东墙 VA、橙虚线 南墙 VA; "
     "(中) ToA — 验证 τ 匹配精度; (下) AoD — 验证 φ 匹配精度。"
     "黑色竖虚线为 LoS/NLoS 分界线 (Y=26.7)。\n"
     "核心结论: 东墙 VA 在 NLoS 区域 27/27 全匹配 (红线连续), 而 argmax (灰虚线) 被穿透径劫持。",
     "论文位置建议: Section IV Results / 主实验结果"),

    ("Fig. 6: VA 匹配单面板增益", "va_matching_result.png",
     "简化版: 仅增益对比, 4 条线。适合正文排版。",
     "论文位置建议: Section IV Results / 增益对比 (与 Fig.5 二选一)"),

    ("Fig. 7: 5 方案波束赋形 SE 对比 (主图)", "isac_beamforming_eval_5schemes.png",
     "1×2 面板: (a) 5 条 SE 曲线 vs UE Y 坐标 — 黑实线 MRT 理想、紫虚线 DFT 码本、"
     "蓝点划线 LoS、橙点线 VA 纯反射、红实线+marker VA 自适应; "
     "(b) 波束训练开销柱状图 — A=64, B=64, C=1, D=1, E=1。\n"
     "核心结论: VA 纯反射 (D) 在 LoS 区域明显低于其他方案 (仅 8.636 vs 13.154 bps/Hz), "
     "但 NLoS 区域与穷举搜索持平 (5.796 bps/Hz)。\n"
     "待改进: 方案 E 加入几何约束后, LoS 区域将对齐 A (13.154), "
     "NLoS 区域退回 D (5.796, 不被穿透径欺骗)。",
     "论文位置建议: Section IV Results / 主对比图"),

    ("Fig. 8: 波束训练开销对比", "isac_beamforming_eval_5schemes.png",
     "同 Fig.7 的右侧面板, 可独立裁切使用。对数坐标, 清晰展示 64 vs 1 的 98.4% 开销降低。",
     "论文位置建议: Section IV Results / 开销分析子图"),
]

for title, filename, description, position in figures:
    doc.add_heading(title, level=2)
    filepath = os.path.join(beam_dir, filename)
    if os.path.exists(filepath):
        if filename.endswith('.png'):
            doc.add_picture(filepath, width=Inches(5.5))
            doc.add_paragraph(f"文件: {filename}").runs[0].font.size = Pt(8)
    else:
        doc.add_paragraph(f"(文件未找到: {filename})")
    p = doc.add_paragraph()
    p.add_run("内容: ").bold = True
    p.add_run(description)
    p2 = doc.add_paragraph()
    p2.add_run(position).italic = True

# ═══════════════════════════════════════════════════════════════════════════
# 4. 关键技术公式
# ═══════════════════════════════════════════════════════════════════════════
doc.add_page_break()
doc.add_heading("4. 关键技术公式", level=1)

formulas = [
    ("虚拟锚点镜像公式 (VAprompt.py)",
     "p' = p - 2 · (n · p + d) · n\n"
     "其中 p = BS 坐标, n = 墙面单位法向量, d = 墙面平面方程常数项"),

    ("3D 线面交点求反射点 (va_matching.py, reflection_point)",
     "t = n · (p0 - VA) / n · (UE - VA)\n"
     "R = VA + t · (UE - VA)\n"
     "其中 p0 = (BS + VA)/2 为墙面中点, direction = UE - VA 为射线方向"),

    ("相位解卷绕角度差 (va_matching.py, phase_diff)",
     "Δφ = arctan2(sin(φ_a - φ_b), cos(φ_a - φ_b)) ∈ [-π, π]\n"
     "避免直接相减在 ±π 边界处的 2π 跳变问题"),

    ("双域匹配判定 (va_matching.py, geometric_match)",
     "match ⇔ |τ_meas - τ_theory| < 1 ns  AND  |Δφ| < 0.1 rad (~5.7°)"),

    ("2D-DFT 码本生成 (beamforming_evaluation_new.py, generate_2d_dft_codebook)",
     "a_y(ky)[n] = exp(j·2π·n·ky / Ny)  for n = 0,...,Ny-1\n"
     "a_z(kz)[m] = exp(j·2π·m·kz / Nz)  for m = 0,...,Nz-1\n"
     "w_{ky,kz} = a_z(kz) ⊗ a_y(ky) / √(Ny·Nz)   ∈ C^{64}"),

    ("全径合成等效信道 (beamforming_evaluation_new.py, compute_h_total)",
     "H_total = Σ_k h_k · exp(-j·2π·fc·τ_k)\n"
     "物理意义: 将所有多径按各自传播相位相干叠加, 得到 BS 天线阵列端口的总等效信道"),

    ("MRT 波束赋形 RSRP (beamforming_evaluation_new.py, compute_mrt_rsrp)",
     "w = h_k^H / ||h_k||\n"
     "h_eq = h_k · w · exp(-j·2π·fc·τ_k)\n"
     "RSRP = |h_eq|² × SNR_SCALE"),

    ("Shannon 频谱效率",
     "SE = log₂(1 + RSRP)  (bps/Hz)"),
]

for name, formula in formulas:
    doc.add_heading(name, level=3)
    p = doc.add_paragraph()
    p.add_run(formula).font.size = Pt(10)

# ═══════════════════════════════════════════════════════════════════════════
# 5. 数据文件清单
# ═══════════════════════════════════════════════════════════════════════════
doc.add_page_break()
doc.add_heading("5. 数据文件清单", level=1)

doc.add_paragraph("输入数据:", style='List Bullet')
doc.add_paragraph("  ue_mimo_multipath_data.npz (607 MB) — MIMO 8×8 多径信道, 122 帧", style='List Bullet')
doc.add_paragraph("  ue_multipath_data.npz (25 MB) — 非 MIMO 标量增益多径, 122 帧", style='List Bullet')
doc.add_paragraph("  isac_pure_sensing.h5 (5.1 MB) — ISAC 感知点云", style='List Bullet')
doc.add_paragraph("  virtual_anchor.json (373 B) — 两个预计算 VA 坐标", style='List Bullet')

doc.add_paragraph("输出图像 (论文可直接使用):", style='List Bullet')
for f in ["view_metis_v35.html (4.8 MB)",
          "wall_extraction_single.html (4.7 MB)",
          "va_visualization.html (4.7 MB)",
          "mimo_dataset_raw_verification.png (1.1 MB)",
          "va_matching_diagnostics.png (664 KB)",
          "va_matching_result.png (283 KB)",
          "isac_beamforming_eval_5schemes.png (516 KB)"]:
    doc.add_paragraph(f"  {f}", style='List Bullet')

# ═══════════════════════════════════════════════════════════════════════════
# 6. 论文写作建议
# ═══════════════════════════════════════════════════════════════════════════
doc.add_heading("6. 论文写作建议", level=1)

doc.add_heading("6.1 建议论文结构", level=2)
sections = [
    "I.    Introduction — ISAC 波束管理动机, 现有方法局限 (穷举开销大, 纯 LoS 不可靠)",
    "II.   System Model — METIS 场景, 8×8 UPA, 信号模型, 传播环境",
    "III.  ISAC-Assisted Beam Tracking — RANSAC 墙面提取 → VA 定位 → 双域匹配 → 自适应选择",
    "IV.   Simulation Results — 5 方案 SE 对比 + 开销分析 + 消融实验",
    "V.    Conclusion"
]
for s in sections:
    doc.add_paragraph(s, style='List Bullet')

doc.add_heading("6.2 拟投稿会议/期刊", level=2)
venues = [
    "IEEE ICC 2027 / GLOBECOM 2026 — ISAC 专题",
    "IEEE Transactions on Wireless Communications",
    "IEEE Wireless Communications Letters (短报文, 4-5 页)",
]
for v in venues:
    doc.add_paragraph(v, style='List Bullet')

doc.add_heading("6.3 当前待完成工作", level=2)
todos = [
    "☐ 方案 E 加入几何约束 (y_coord < LOS_BLOCKAGE_Y), 完成'VA 鉴别穿透伪径'的完整实现",
    "☐ 新增场景几何示意图 (2D 俯视图, matplotlib 手绘), 作为 Fig.1",
    "☐ 新增方案 E 自适应决策流程图, 放入方法小节",
    "☐ DFT 码本量化损耗子图 (ΔSE = MRT - DFT 随帧变化)",
    "☐ 从 HTML 截取高质量 3D 视图截图用于论文 (view_metis_v35, wall_extraction_single, va_visualization)",
    "☐ 撰写正文 (LaTeX), 整合所有图表",
]
for t in todos:
    doc.add_paragraph(t, style='List Bullet')

# ── 保存 ──────────────────────────────────────────────────────────────────
output_path = "/home/xu/xdh/isac/ISAC_beam/ISAC_研究进展汇总报告.docx"
doc.save(output_path)
print(f"报告已保存: {output_path}")
